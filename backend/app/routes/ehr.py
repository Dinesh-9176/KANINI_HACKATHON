"""EHR document parsing endpoint."""

import re
from fastapi import APIRouter, UploadFile, File, HTTPException
from docx import Document
import io

router = APIRouter()


def celsius_to_fahrenheit(temp_c: float) -> float:
    """Convert Celsius to Fahrenheit."""
    return temp_c * 9 / 5 + 32


def parse_ehr_docx(file_bytes: bytes) -> dict:
    """Parse a .docx EHR document and extract structured patient data.

    Expected format (from sample Patient_Record_1.docx):
    - Title paragraph: "Patient Record #N"
    - Table 0 (Demographics): Age, Gender rows
    - Table 1 (Vitals): BP Systolic, BP Diastolic, Heart Rate, Temperature, SpO2, Respiratory Rate
    - Table 2 (Clinical): Symptoms (comma-separated), Conditions (comma-separated)
    - Patient Notes paragraph after "Patient Notes" heading
    """
    doc = Document(io.BytesIO(file_bytes))

    result = {
        "name": "",
        "age": 0,
        "gender": "",
        "blood_pressure_systolic": None,
        "blood_pressure_diastolic": None,
        "heart_rate": None,
        "temperature": None,
        "oxygen_saturation": None,
        "respiratory_rate": None,
        "symptoms": [],
        "conditions": [],
        "notes": "",
    }

    # Extract patient name from title if present
    for para in doc.paragraphs:
        if para.style.name == "Title" and para.text.strip():
            result["name"] = para.text.strip()
            break

    # Extract notes from paragraphs after "Patient Notes" heading
    capture_notes = False
    for para in doc.paragraphs:
        if "patient notes" in para.text.lower() and para.style.name.startswith("Heading"):
            capture_notes = True
            continue
        if capture_notes and para.text.strip():
            result["notes"] = para.text.strip()
            capture_notes = False

    # Parse tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2:
                continue

            label = cells[0].lower()
            value = cells[1].strip()

            # Demographics
            if label == "age":
                try:
                    result["age"] = int(value)
                except ValueError:
                    pass
            elif label == "gender":
                result["gender"] = value.lower()

            # Vitals
            elif "systolic" in label:
                try:
                    result["blood_pressure_systolic"] = int(float(value))
                except ValueError:
                    pass
            elif "diastolic" in label:
                try:
                    result["blood_pressure_diastolic"] = int(float(value))
                except ValueError:
                    pass
            elif "heart rate" in label:
                try:
                    result["heart_rate"] = int(float(value))
                except ValueError:
                    pass
            elif "temperature" in label:
                try:
                    temp = float(value)
                    # Check unit - if in Celsius (< 50), convert to Fahrenheit
                    unit = cells[2].strip().upper() if len(cells) > 2 else ""
                    if "C" in unit or temp < 50:
                        temp = celsius_to_fahrenheit(temp)
                    result["temperature"] = round(temp, 1)
                except (ValueError, IndexError):
                    pass
            elif "oxygen" in label or "spo2" in label:
                try:
                    result["oxygen_saturation"] = int(float(value))
                except ValueError:
                    pass
            elif "respiratory" in label:
                try:
                    result["respiratory_rate"] = int(float(value))
                except ValueError:
                    pass

            # Clinical
            elif "symptom" in label:
                symptoms = [s.strip() for s in value.split(",") if s.strip()]
                result["symptoms"] = symptoms
            elif "condition" in label:
                conditions = [c.strip() for c in value.split(",") if c.strip()]
                result["conditions"] = conditions
            elif "name" in label and "patient" in label:
                result["name"] = value

    return result


@router.post("/parse-ehr")
async def parse_ehr(file: UploadFile = File(...)):
    """Parse an uploaded EHR document (.docx) and return structured patient data."""

    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="Only .docx files are supported. Please upload a Word document.",
        )

    try:
        file_bytes = await file.read()
        parsed_data = parse_ehr_docx(file_bytes)
        return parsed_data
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail=f"Failed to parse EHR document: {str(e)}",
        )
